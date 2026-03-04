"""Lexio Pipeline — DOCX generator (parameterized from config)."""

import logging
import os
import uuid
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from packages.pipeline.pipeline_config import PipelineConfig

logger = logging.getLogger("lexio.pipeline.docx")

OUTPUT_DIR = Path("/app/output") if os.path.exists("/app") else Path("output")


async def generate_docx(
    text: str,
    context: dict,
    config: PipelineConfig,
) -> str:
    """Generate a DOCX file from the final text."""

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = DocxDocument()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    # Default styles
    font_name = context.get("font_name", "Times New Roman")
    font_size = int(context.get("font_size", 12))
    line_spacing = float(context.get("line_spacing", 1.5))

    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size)
    style.paragraph_format.line_spacing = line_spacing

    # Split text into paragraphs
    paragraphs = text.split("\n\n")

    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue

        # Detect section headers (ALL CAPS lines)
        is_header = (
            para_text == para_text.upper()
            and len(para_text) < 200
            and not para_text.startswith("[")
        )

        if is_header:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(para_text)
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(font_size)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(para_text)
            run.font.name = font_name
            run.font.size = Pt(font_size)
            p.paragraph_format.first_line_indent = Cm(1.25)
            p.paragraph_format.space_after = Pt(6)

    # Save
    doc_type = config.document_type_id
    filename = f"{doc_type}_{context.get('document_id', uuid.uuid4().hex[:8])}.docx"
    filepath = OUTPUT_DIR / filename
    doc.save(str(filepath))

    size_kb = filepath.stat().st_size / 1024
    logger.info(f"DOCX generated: {filepath} ({size_kb:.1f} KB)")

    return str(filepath)
